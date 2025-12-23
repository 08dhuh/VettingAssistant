import os
from typing import Optional, Union
import openai
from docx import Document
import PyPDF2
import io
from ..models import DatasetMetadata


class DescriptionParser:
    """
    Parses description documents using OpenAI to extract structured metadata.
    Supports TXT, DOCX, and PDF formats.
    """

    def __init__(self, api_key: Optional[str] = None, model: str = "gpt-4o"):
        """
        Initialize the parser.

        Args:
            api_key: OpenAI API key (defaults to OPENAI_API_KEY env var)
            model: OpenAI model to use for parsing
        """
        self.api_key = api_key or os.environ.get("OPENAI_API_KEY")
        self.model = model

        if not self.api_key:
            raise ValueError("OpenAI API key not provided and not found in environment")

        self.client = openai.OpenAI(api_key=self.api_key)

    def parse(self, file_content: Union[str, bytes, io.BytesIO], file_type: str = "txt") -> DatasetMetadata:
        """
        Parse description document and extract metadata.

        Args:
            file_content: File content (string for txt, bytes/BytesIO for docx/pdf)
            file_type: File type ('txt', 'docx', 'pdf')

        Returns:
            DatasetMetadata object with extracted information
        """
        # Extract text from file
        text_content = self._extract_text(file_content, file_type)

        # Use OpenAI to parse the description
        metadata = self._parse_with_openai(text_content)

        return metadata

    def _extract_text(self, file_content: Union[str, bytes, io.BytesIO], file_type: str) -> str:
        """Extract text from various file formats"""
        if file_type == "txt" or file_type == "md":
            if isinstance(file_content, str):
                return file_content
            elif isinstance(file_content, bytes):
                return file_content.decode("utf-8")
            elif isinstance(file_content, io.BytesIO):
                return file_content.read().decode("utf-8")

        elif file_type == "docx":
            if isinstance(file_content, io.BytesIO):
                doc = Document(file_content)
            else:
                doc = Document(io.BytesIO(file_content))

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text

        elif file_type == "pdf":
            if isinstance(file_content, io.BytesIO):
                pdf_reader = PyPDF2.PdfReader(file_content)
            else:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _parse_with_openai(self, text_content: str) -> DatasetMetadata:
        """Use OpenAI to extract structured metadata from description text"""

        prompt = f"""
You are an expert data analyst. Extract structured metadata from the following data description document.

IMPORTANT: Return ONLY valid JSON with no additional text, markdown, or explanations.

Extract the following information:
1. **population**: The population used to derive the results (string)
2. **method_of_analysis**: The method of analysis used (string)
3. **datasets_used**: List of datasets used (array of strings)
4. **data_description**: Overall description of the data (string)
5. **variables**: Dictionary mapping variable/column names to their descriptions (object)
6. **percentage_columns**: List of column names that represent percentages, shares, or proportions (array of strings)
7. **sample_size_column**: Name of the column containing sample sizes/counts, if mentioned (string or null)

For REGRESSION outputs, also extract:
8. **regression_type**: Type of regression model (string or null)
   - If mentions "linear regression", "OLS", or just "regression" → "OLS"
   - If mentions "logit" or "logistic" → "logit"
   - If mentions "probit" → "probit"
   - If not a regression output → null

9. **variable_types**: For each variable, infer its type based on description (object)
   - If description contains "indicator", "dummy", "binary", "flag", "yes/no", "0/1" → "binary"
   - If description contains "income", "age", "amount", "value", "price", "salary", "wage", "count", "number of", "rate", "score", "index" → "continuous"
   - If unclear or no keywords match → "unknown"

10. **outcome_variable**: What the regression is predicting/modeling (string or null)
    - Extract the dependent variable or outcome being modeled
    - Keep it concise (e.g., "number of children", "employment status")
    - null if not a regression output

Description Document:
{text_content}

Return JSON in this exact format:
{{
  "population": "string or null",
  "method_of_analysis": "string or null",
  "datasets_used": ["dataset1", "dataset2"],
  "data_description": "string or null",
  "variables": {{
    "variable_name": "description",
    "another_variable": "description"
  }},
  "percentage_columns": ["column1", "column2"],
  "sample_size_column": "column_name or null",
  "regression_type": "OLS or logit or probit or null",
  "variable_types": {{
    "variable_name": "continuous or binary or unknown"
  }},
  "outcome_variable": "string or null"
}}

If information is not available, use null for strings or [] for arrays or {{}} for objects.
"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are a data analyst expert. Extract structured metadata from data descriptions. Return ONLY valid JSON with no markdown formatting."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,
                max_tokens=2000,
            )

            # Parse the response
            response_text = response.choices[0].message.content.strip()

            # Remove markdown code blocks if present
            if response_text.startswith("```json"):
                response_text = response_text[7:]
            if response_text.startswith("```"):
                response_text = response_text[3:]
            if response_text.endswith("```"):
                response_text = response_text[:-3]

            response_text = response_text.strip()

            # Parse JSON
            import json
            parsed_data = json.loads(response_text)

            # Create DatasetMetadata object
            metadata = DatasetMetadata(
                population=parsed_data.get("population"),
                method_of_analysis=parsed_data.get("method_of_analysis"),
                datasets_used=parsed_data.get("datasets_used", []),
                data_description=parsed_data.get("data_description"),
                variables=parsed_data.get("variables", {}),
                percentage_columns=parsed_data.get("percentage_columns", []),
                sample_size_column=parsed_data.get("sample_size_column"),
                raw_description=text_content,
                regression_type=parsed_data.get("regression_type"),
                variable_types=parsed_data.get("variable_types", {}),
                outcome_variable=parsed_data.get("outcome_variable"),
            )

            return metadata

        except Exception as e:
            # Return minimal metadata if parsing fails
            return DatasetMetadata(
                raw_description=text_content,
            )